import streamlit as st
import os
import time
import re
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from google import genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from io import BytesIO

# 1. Buka kunci (.env)
load_dotenv(override=True)
try:
    api_key_google = st.secrets["GOOGLE_API_KEY"] # Untuk Streamlit Cloud
except:
    api_key_google = os.getenv("GOOGLE_API_KEY")  # Untuk Laptop Lokal

os.environ["GOOGLE_API_KEY"] = api_key_google

# --- FUNGSI AUTO-RETRY ---
def cuba_jana_ai(prompt_teks):
    client = genai.Client(api_key=api_key_google)
    senarai_model = ["gemini-flash-latest", "gemini-3.5-flash", "gemini-3.1-flash-lite"] 
    max_cuba = 3 
    for cubaan in range(max_cuba):
        for model_ai in senarai_model:
            try:
                respons = client.models.generate_content(model=model_ai, contents=prompt_teks)
                return respons
            except Exception as e:
                ralat = str(e)
                if "503" in ralat or "429" in ralat:
                    time.sleep(5) 
                    continue 
                elif "404" in ralat:
                    continue 
                else:
                    raise e 
    raise Exception("Pelayan Google sesak. Sila cuba lagi.")

# --- FUNGSI BINA FAIL WORD AP & PENGURUSAN ---
def bina_fail_word(teks_ai, tajuk_dokumen, metadata, is_pengurusan=False, is_arahan_amalan=False):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12) # Saiz 12 standard AP
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(12)

    # --- 1. PEMBINAAN KEPALA SURAT ---
    if not is_pengurusan:
        # Format Khas Alasan Penghakiman (AP)
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        run_p1 = p1.add_run(metadata['pihak1'].upper())
        run_p1.bold = True

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run_lwn = p2.add_run("lwn")
        run_lwn.italic = True

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(24)
        run_p3 = p3.add_run(metadata['pihak2'].upper())
        run_p3.bold = True

        p4 = doc.add_paragraph(f"[{metadata['mahkamah']} {metadata['negeri']}, {metadata['hakim']}, pada {metadata['tarikh']}]")
        p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p4.paragraph_format.space_after = Pt(0)
        
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p5.paragraph_format.space_after = Pt(12)
        p5.add_run(f"[Kes No. {metadata['nokes']}]\n[{metadata['jenis_p']}]")

        if metadata['peguam']:
            p6 = doc.add_paragraph(f"Peguam Syarie: {metadata['peguam']}")
            p6.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p6.paragraph_format.space_after = Pt(24)
    else:
        # Format Khas Kertas Kerja / Kertas Konsep
        p_tajuk = doc.add_paragraph()
        p_tajuk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tajuk.paragraph_format.space_after = Pt(24)
        run_t = p_tajuk.add_run(tajuk_dokumen.upper())
        run_t.bold = True

    # --- 2. PEMBERSIHAN SAMPAH MARKDOWN AI & HALUSINASI ---
    teks_bersih = teks_ai.replace('**', '').replace('##', '').replace('#', '').replace('*', '').replace('>', '')
    lines = teks_bersih.split('\n')
    
    senarai_hitam_header = ["DI DALAM MAHKAMAH", "DALAM NEGERI", "ANTARA", "DENGAN", "DI HADAPAN", "PERMOHONAN NO", "ALASAN PENGHAKIMAN"]
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if any(line.upper().startswith(haram) for haram in senarai_hitam_header) and not is_pengurusan:
            i += 1
            continue

        # FORMAT: Tajuk KEPUTUSAN (Center, Bold)
        if line.upper() == "KEPUTUSAN" or line.upper() == "5. KEPUTUSAN" or line.upper() == "4. KEPUTUSAN":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run("KEPUTUSAN")
            run.bold = True
            
        # FORMAT: Ayat "SETELAH..." (Bold & Underline perkataan pertama)
        elif line.upper().startswith("SETELAH") and not is_pengurusan:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            words = line.split(" ", 1)
            run_setelah = p.add_run(words[0].upper() + " ")
            run_setelah.bold = True
            run_setelah.underline = True
            if len(words) > 1:
                p.add_run(words[1])

        # FORMAT: Tajuk Utama (Kiri, Bold) - Termasuk tajuk Arahan Amalan
        elif line.upper() in ["PERMOHONAN", "FAKTA KES", "ULASAN MAHKAMAH", "UNDANG-UNDANG YANG DIPAKAI"] or re.match(r'^\d+\.\s+[A-Z\s]+$', line.upper()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run(line.upper())
            run.bold = True
            
        # KESAN JADUAL (Pengecualian untuk Mod 2 - Arahan Amalan TIDAK perlukan jadual)
        elif line.startswith('|') and not is_arahan_amalan:
            table_data = []
            while i < len(lines) and line.startswith('|'):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                if cells: table_data.append(cells)
                i += 1
                if i < len(lines): line = lines[i].strip()
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row_cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text.replace('<br>', '\n')
            continue 

        # FORMAT: Teks Biasa (Justify)
        else:
            line = re.sub(r'^\[\d+\]\s*', '', line)
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI STREAMLIT ---
st.set_page_config(page_title="Pintar Syariah AI", page_icon="⚖️", layout="wide")
st.title("⚖️ Artificial Intelligence Mahkamah Syariah (AIMS)")
st.divider()

tab_kes, tab_pengurusan = st.tabs(["🏛️ Mod 1: Analisis Kes Syariah (AP)", "📝 Mod 2: Kertas Kerja & Arahan Amalan"])

# ==========================================
# MOD 1: ANALISIS KES (DRAF AP PRO)
# ==========================================
with tab_kes:
    col_meta, col_input = st.columns([1, 2]) 
    
    with col_meta:
        st.subheader("📋 Maklumat Kes (Metadata)")
        m_negeri = st.selectbox("Bidang Kuasa (Negeri):", ["Selangor", "Wilayah Persekutuan", "Johor", "Perak", "Kedah", "Kelantan", "Terengganu", "Pahang", "Negeri Sembilan", "Melaka", "Pulau Pinang", "Perlis", "Sabah", "Sarawak"])
        m_level = st.selectbox("Hierarki Mahkamah:", ["Mahkamah Rendah Syariah", "Mahkamah Tinggi Syariah", "Mahkamah Rayuan Syariah"])
        m_hakim = st.text_input("Nama Hakim / Panel:", placeholder="Cth: YA Tuan Haji...")
        m_tarikh = st.text_input("Tarikh Sidang / Keputusan:", placeholder="Cth: 26 Ramadhan 1438H / 21 Jun 2017")
        m_nokes = st.text_input("No. Kes:", placeholder="Cth: 10000-003-0001-2024")
        m_pihak1 = st.text_area("Pihak 1 (Plaintif/Pemohon):", height=70)
        m_pihak2 = st.text_area("Pihak 2 (Defendan/Responden):", height=70)
        m_jenis = st.text_area("Jenis Permohonan (Peruntukan):", placeholder="Cth: Permohonan Semakan di bawah seksyen 68...", height=70)
        m_peguam = st.text_input("Nama Peguam (Jika ada):", placeholder="Cth: Mohd. Faiz b. Adnan...")

    with col_input:
        st.subheader("📝 Ringkasan Fakta & Seksyen Utama")
        f_permohonan = st.text_area("Butiran PERMOHONAN (Jika ada):", placeholder="Biarkan kosong jika mahu AI drafkan...")
        f_fakta = st.text_area("Butiran FAKTA KES (Wajib isi):", height=150)
        f_ulasan = st.text_area("Butiran ULASAN MAHKAMAH (Jika ada):")
        f_keputusan = st.text_area("Butiran KEPUTUSAN (Jika ada):")

        if st.button("🔍 Jana Draf Alasan Penghakiman (AP)", type="primary"):
            if f_fakta.strip() == "":
                st.warning("⚠️ Sila masukkan Fakta Kes!")
            else:
                with st.spinner("AI sedang merangka AP mengikut format Mahkamah..."):
                    try:
                        embeddings = GoogleGenerativeAIEmbeddings(model="gemini-embedding-001")
                        db = Chroma(persist_directory="./database_vektor_google", embedding_function=embeddings)
                        retriever = db.as_retriever(search_kwargs={"k": 5})
                        dokumen_relevan = retriever.invoke(f_fakta)
                        konteks = "\n".join([d.page_content for d in dokumen_relevan])

                        prompt_ap = f"""Anda adalah Penyelidik Undang-Undang Kanan. Tugas anda merangka kandungan ALASAN PENGHAKIMAN (AP).

AMARAN KERAS FORMAT (WAJIB PATUH 100%):
1. JANGAN tulis Kepala Surat sama sekali.
2. Mula draf anda TERUS dengan tajuk: PERMOHONAN
3. DILARANG menggunakan simbol Markdown (*, >, #, ##).
4. DILARANG menggunakan kurungan bernombor seperti [1], [2], [3] pada permulaan perenggan.
5. DILARANG menggunakan HURUF BESAR (ALL CAPS) untuk perenggan huraian. 
6. JANGAN SESEKALI menyalin maklumat pihak, nama hakim, tarikh, atau no kes daripada Kes Rujukan.

STRUKTUR KANDUNGAN:
PERMOHONAN
(Tulis draf permohonan. Ringkasan: {f_permohonan})

FAKTA KES
(Huraikan fakta. Input: {f_fakta})

ULASAN MAHKAMAH
(Salin PERUNTUKAN UNDANG-UNDANG verbatim dari rujukan kes bagi negeri {m_negeri}. Ulas menggunakan ayat biasa. Ringkasan: {f_ulasan})

KEPUTUSAN
(Mesti dimulakan dengan ayat: "SETELAH Kami membaca dan meneliti permohonan..." ATAU "SETELAH Mahkamah meneliti...")
(Guna laras bahasa hierarki {m_level}. Ringkasan: {f_keputusan})

RUJUKAN KES LEPAS:
{konteks}
"""
                        respons = cuba_jana_ai(prompt_ap)
                        
                        meta_dict = {
                            'negeri': m_negeri, 'mahkamah': m_level, 'hakim': m_hakim, 'tarikh': m_tarikh, 
                            'nokes': m_nokes, 'pihak1': m_pihak1, 'pihak2': m_pihak2, 
                            'jenis_p': m_jenis, 'peguam': m_peguam
                        }
                        
                        fail_word = bina_fail_word(respons.text, "ALASAN PENGHAKIMAN", meta_dict)
                        
                        st.success("✅ Draf selesai dijana! Sila muat turun fail Word di bawah.")
                        st.download_button("📄 Muat Turun AP (Word)", data=fail_word, file_name=f"AP_{m_nokes}.docx")
                    except Exception as e: st.error(f"Ralat: {e}")

# ==========================================
# MOD 2: KERTAS KERJA PENGURUSAN & ARAHAN AMALAN
# ==========================================
with tab_pengurusan:
    st.info("Kertas Kerja Pengurusan & Kertas Konsep Arahan Amalan")
    jenis_kertas = st.selectbox("Jenis Dokumen:", ["Kertas Kerja Bengkel / Program", "Kertas Kerja Bajet", "Kertas Konsep Arahan Amalan"])
    
    # BORANG PINTAR: Borang berubah mengikut pilihan Jenis Kertas
    if jenis_kertas == "Kertas Konsep Arahan Amalan":
        bahagian_unit = st.text_input("1. Bahagian / Unit Penyedia:", value="Bahagian Dasar & Penyelidikan (BPKR)")
        nama_program = st.text_input("2. Tajuk Kertas Konsep:", placeholder="Cth: Penangguhan Kes Atas Alasan Sijil Cuti Sakit")
        
        col1, col2 = st.columns(2)
        with col1:
            latar_belakang = st.text_area("3. Latar Belakang / Ringkasan Isu:", height=150)
        with col2:
            asas_pertimbangan = st.text_area("4. Asas Pertimbangan (Hujah Hukum Syarak/Undang-undang):", height=150)
            
        cadangan_syor = st.text_area("5. Cadangan & Syor:", height=100)
        
    else:
        # Borang Lama Untuk Bajet & Program
        bahagian_unit = st.text_input("1. Bahagian / Unit Penyedia:")
        nama_program = st.text_input("2. Nama Program / Aktiviti:")
        tarikh_masa = st.text_input("3. Tarikh, Masa & Tempoh:")
        tempat_program = st.text_input("4. Tempat / Lokasi Program:")
        
        col1, col2 = st.columns(2)
        with col1:
            maklumat_peserta = st.text_area("5. Maklumat & Bilangan Peserta:", height=150)
            maklumat_penceramah = st.text_area("6. Maklumat Penceramah:", height=150)
        with col2:
            kos_makan_minum = st.text_area("7. Anggaran Kos Kewangan:", height=150)
            objektif_tambahan = st.text_area("8. Objektif & Latar Belakang Tambahan:", height=150)

    if st.button("🚀 Jana Kertas Cadangan Lengkap", type="primary", key="btn_pengurusan"):
        if nama_program.strip() == "" or bahagian_unit.strip() == "":
            st.warning("⚠️ Sila isi Bahagian/Unit dan Tajuk terlebih dahulu bro!")
        else:
            with st.spinner(f"Menyusun format {jenis_kertas} rasmi..."):
                try:
                    embeddings = GoogleGenerativeAIEmbeddings(model="gemini-embedding-001")
                    db = Chroma(persist_directory="./database_vektor_google", embedding_function=embeddings)
                    
                    # Semak sumber carian vektor yang tepat
                    filter_kategori = "Arahan Amalan" if jenis_kertas == "Kertas Konsep Arahan Amalan" else "Pengurusan"
                    retriever = db.as_retriever(search_kwargs={"k": 5, "filter": {"sumber": filter_kategori}})
                    dokumen_relevan = retriever.invoke(nama_program)
                    konteks_teks = "\n".join([f"\n--- CONTOH TEMPLAT {idx+1} ---\n{doc.page_content}\n" for idx, doc in enumerate(dokumen_relevan)])

                    # PROMPT PINTAR: Berubah mengikut jenis borang
                    if jenis_kertas == "Kertas Konsep Arahan Amalan":
                        prompt_pengurusan = f"""Anda adalah Penyelidik Kanan Jabatan Kehakiman Syariah Malaysia. Bina draf Kertas Konsep Arahan Amalan yang rasmi.

STRUKTUR WAJIB (JANGAN GUNA JADUAL & DILARANG GUNA SIMBOL BINTANG **):
JANGAN tulis tajuk utama, terus mula dengan format nombor di bawah:

1. TUJUAN 
(Jelaskan tujuan kertas konsep ini disediakan untuk pertimbangan YAA Ketua Hakim Syarie)

2. LATAR BELAKANG
(Huraikan isu/kronologi. Input: {latar_belakang})

3. ASAS-ASAS PERTIMBANGAN
(Berikan hujah undang-undang/hukum syarak. Input: {asas_pertimbangan})

4. CADANGAN DAN SYOR
(Input: {cadangan_syor})

5. CONTOH ARAHAN AMALAN
(Bina SATU draf contoh Arahan Amalan yang lengkap dan kemas bermula dengan 'ARAHAN AMALAN NO. [...] TAHUN [...]')

Rujuk gaya bahasa rasmi dari dokumen ini:
{konteks_teks}
"""
                    else:
                        prompt_pengurusan = f"""Anda adalah Pegawai Tadbir Kanan di JKSM. Jana draf {jenis_kertas} yang rasmi.

PANDUAN STRUKTUR TEKS:
- DILARANG menggunakan simbol bintang (**) untuk tulisan Bold.
- JANGAN tulis tajuk utama. Terus mula dengan nombor 1 di bawah:

1. BAHAGIAN/UNIT 
{bahagian_unit}
2. TUJUAN 
3. LATAR BELAKANG ({objektif_tambahan})
4. OBJEKTIF PROGRAM 
5. BUTIR-BUTIR PROGRAM
(Tarikh: {tarikh_masa}, Tempat: {tempat_program}, Peserta: {maklumat_peserta}, Penceramah: {maklumat_penceramah})
6. ANGGARAN PERBELANJAAN
7. KESIMPULAN / PENUTUP

INPUT DATA KEWANGAN:
{kos_makan_minum}

ARAHAN KRITIKAL PEMBINAAN JADUAL (AMARAN KERAS):
1. Anda MESTI membina SATU JADUAL SAHAJA di bawah bahagian 6. ANGGARAN PERBELANJAAN.
2. Jadual tersebut WAJIB mempunyai TEPAT 3 KOLUM: | Bil. | Butiran | Jumlah |
3. Masukkan SEMUA butiran pengiraan ke dalam kolum "Butiran". Gunakan <br> untuk baris baharu di dalam sel.
4. Baris terakhir jadual mesti bernama JUMLAH KESELURUHAN.

Rujuk gaya bahasa dokumen ini:
{konteks_teks}
"""
                    respons = cuba_jana_ai(prompt_pengurusan)
                    st.divider()
                    st.subheader(f"💡 Hasil Penjanaan AI ({jenis_kertas})")
                    st.write(respons.text)

                    meta_dummy = {'mahkamah':'', 'hakim':'', 'tarikh':'', 'nokes':'', 'pihak1':'', 'pihak2':'', 'jenis_p':'', 'peguam':'', 'negeri':''}
                    
                    # LOGIK TAJUK DOKUMEN DALAM WORD
                    is_aa = (jenis_kertas == "Kertas Konsep Arahan Amalan")
                    tajuk_rasmi = f"KERTAS KONSEP ARAHAN AMALAN\n{nama_program.upper()}" if is_aa else f"KERTAS PERMOHONAN KELULUSAN BERBELANJA BAGI\n{nama_program.upper()}\nJABATAN KEHAKIMAN SYARIAH MALAYSIA"
                    
                    fail_pengurusan_docx = bina_fail_word(respons.text, tajuk_rasmi, meta_dummy, is_pengurusan=True, is_arahan_amalan=is_aa)
                    st.download_button(f"📄 Muat Turun {jenis_kertas} (Word)", data=fail_pengurusan_docx, file_name=f"{jenis_kertas.replace(' ', '_')}.docx")

                except Exception as e: st.error(f"❌ Ralat Sistem: {e}")

# ==========================================
# FOOTER HAK CIPTA (HAK MILIK EKSKLUSIF)
# ==========================================
footer_html = """
<style>
.footer {
    position: fixed;
    left: 0;
    bottom: 0;
    width: 100%;
    background-color: #0F172A; /* Warna latar belakang gelap */
    color: #94A3B8; /* Warna tulisan kelabu cair */
    text-align: center;
    padding: 12px;
    font-size: 13px;
    border-top: 4px solid #1D4ED8; /* Garisan biru kat atas footer */
    z-index: 999;
}
.block-container { padding-bottom: 80px; }
</style>
<div class="footer">
    © 2026 AIMS Dibangunkan Oleh Pihak UTES. Hak Cipta Terpelihara.
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)